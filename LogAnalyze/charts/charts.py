# encoding:utf-8
"""
生成图表
============
Date@2019/05/21
Author@wangjunxiong
"""
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


class generate_charts(object):
    """generate charts and final file"""
    def __init__(self):
        """
        :param src_file:
        :return charts and final file
        """
        pass

    def generate_data(self):
        """ 采用 {k, v} 进行关键数据的统计并存储
        """
        print "minutes文件处理中..."
        with open("rst_m.txt", 'r')as f:
            self.m_count_dic, self.m_info_dic, self.m_warn_dic = {}, {}, {}  # 分别记录分钟数量，info数量，出错数量
            self.m_x, self.h_x, self.d_x = [], [], []
            lines = f.readlines()
            for line in lines:
                line = [int(i) for i in line.split()]
                tmp_m, tmp_info, tmp_warn = line
                if tmp_m not in self.m_x:
                    self.m_x.append(tmp_m)
                # 记录分钟对应info数量
                if self.m_info_dic.has_key(tmp_m):
                    self.m_info_dic[tmp_m] += tmp_info
                else:
                    self.m_info_dic[tmp_m] = tmp_info
                # 记录分钟对应warn数量
                if self.m_warn_dic.has_key(tmp_m):
                    self.m_warn_dic[tmp_m] += tmp_warn
                else:
                    self.m_warn_dic[tmp_m] = tmp_warn
                # 记录分钟数量，求平均值
                if self.m_count_dic.has_key(tmp_m):
                    self.m_count_dic[tmp_m] += 1
                else:
                    self.m_count_dic[tmp_m] = 1

        print "hours文件处理中..."
        with open("rst_h.txt", 'r')as f:
            self.h_count_dic, self.h_info_dic, self.h_warn_dic = {}, {}, {}  # 分别记录分钟数量，info数量，出错数量
            lines = f.readlines()
            for line in lines:
                line = [int(i) for i in line.split()]
                tmp_h, tmp_info, tmp_warn = line
                if tmp_h not in self.h_x:
                    self.h_x.append(tmp_h)
                # 记录分钟对应info数量
                if self.h_info_dic.has_key(tmp_h):
                    self.h_info_dic[tmp_h] += tmp_info
                else:
                    self.h_info_dic[tmp_h] = tmp_info
                # 记录分钟对应warn数量
                if self.h_warn_dic.has_key(tmp_h):
                    self.h_warn_dic[tmp_h] += tmp_warn
                else:
                    self.h_warn_dic[tmp_h] = tmp_warn
                # 记录分钟数量，求平均值
                if self.h_count_dic.has_key(tmp_h):
                    self.h_count_dic[tmp_h] += 1
                else:
                    self.h_count_dic[tmp_h] = 1

        print "days文件处理中..."
        with open("rst_d.txt", 'r')as f:
            self.d_count_dic, self.d_info_dic, self.d_warn_dic = {}, {}, {}  # 分别记录分钟数量，info数量，出错数量
            lines = f.readlines()
            for line in lines:
                line = [int(i) for i in line.split()]
                tmp_d, tmp_info, tmp_warn = line
                if tmp_d not in self.d_x:
                    self.d_x.append(tmp_d)
                # 记录分钟对应info数量
                if self.d_info_dic.has_key(tmp_d):
                    self.d_info_dic[tmp_d] += tmp_info
                else:
                    self.d_info_dic[tmp_d] = tmp_info
                # 记录分钟对应warn数量
                if self.d_warn_dic.has_key(tmp_d):
                    self.d_warn_dic[tmp_d] += tmp_warn
                else:
                    self.d_warn_dic[tmp_d] = tmp_warn
                # 记录分钟数量，求平均值
                if self.d_count_dic.has_key(tmp_d):
                    self.d_count_dic[tmp_d] += 1
                else:
                    self.d_count_dic[tmp_d] = 1
        self.generate_m_charts()
        self.generate_h_charts()
        self.generate_d_charts()
        self.fuck_end()
        pass

    def generate_m_charts(self):
        print "minutes图表生成..."
        # print self.m_info_dic
        # print self.m_warn_dic
        # print self.m_count_dic
        x = [int(i) for i in self.m_count_dic]
        # x.sort()
        info_y, warn_y, total_y, total_y, right_persent_y = [], [], [], [], []
        for i in x:
            info_y.append(int(self.m_info_dic[i]) / int(self.m_count_dic[i]))  # 分钟平均探测INFO数
            warn_y.append(int(self.m_warn_dic[i]) / int(self.m_count_dic[i]))  # 分钟平均探测WARN数
            total_y.append(int(self.m_info_dic[i]) + int(self.m_warn_dic[i]))  # 分钟评卷探测 总数
        for i in x:
            # 探测正确率
            right_persent_y.append(float(self.m_info_dic[i]) / float(int(self.m_info_dic[i]) + int(self.m_warn_dic[i])))
        x = self.m_x
        # plt.bar(x, info_y, width=0.5, color='b', lw=100)
        # plt.grid(True)
        # plt.legend(loc=0)
        # plt.axis('tight')
        # plt.xlabel('minutes')
        # plt.title("success / minutes (average)")
        # plt.savefig("charts/image/m_info.jpg")

        plt.plot(x, info_y)
        plt.title("success / minutes (average)")
        plt.savefig("charts/image/m_info.jpg")
        plt.close()

        plt.plot(x, warn_y)
        plt.title("warn / minutes (average)")
        plt.savefig("charts/image/m_warn.jpg")
        plt.close()

        plt.plot(x, total_y)
        plt.title("total / minutes")
        plt.savefig("charts/image/m_total.jpg")
        plt.close()

        plt.plot(x, right_persent_y)
        plt.title("accuracy rate")
        plt.savefig("charts/image/m_rightP.jpg")
        plt.close()
        pass

    def generate_h_charts(self):
        print "hours图表生成..."
        x = [int(i) for i in self.h_count_dic]
        # print x
        # x.sort()
        info_y, warn_y, total_y, total_y, right_persent_y = [], [], [], [], []
        for i in x:
            info_y.append(int(self.h_info_dic[i]) / int(self.h_count_dic[i]))  # 分钟平均探测INFO数
            warn_y.append(int(self.h_warn_dic[i]) / int(self.h_count_dic[i]))  # 分钟平均探测WARN数
            total_y.append(int(self.h_info_dic[i]) + int(self.h_warn_dic[i]))  # 分钟评卷探测 总数
        for i in x:
            # 探测正确率
            right_persent_y.append(float(self.h_info_dic[i]) / float(int(self.h_info_dic[i]) + int(self.h_warn_dic[i])))
        x = self.h_x
        plt.bar(x, info_y, width=0.7, color='g', lw=100)
        plt.grid(True)
        plt.legend(loc=0)
        plt.axis('tight')
        plt.xlabel('hour')
        plt.title("success / hour (average)")

        # plt.plot(x, info_y)
        # plt.title("success / hour (average)")
        plt.savefig("charts/image/h_info.jpg")
        plt.close()

        plt.plot(x, warn_y)
        plt.title("warn / hour")
        plt.savefig("charts/image/h_warn.jpg")
        plt.close()

        plt.plot(x, total_y)
        plt.title("total / hour")
        plt.savefig("charts/image/h_total.jpg")
        plt.close()

        plt.plot(x, right_persent_y)
        plt.title("accuracy rate")
        plt.savefig("charts/image/h_rightP.jpg")
        plt.close()
        pass

    def generate_d_charts(self):
        print "days图表生成..."
        x = [int(i) for i in self.d_count_dic]
        # x.sort()
        info_y, warn_y, total_y, total_y, right_persent_y = [], [], [], [], []
        for i in x:
            info_y.append(int(self.d_info_dic[i]) / int(self.d_count_dic[i]))  # 分钟平均探测INFO数
            warn_y.append(int(self.d_warn_dic[i]) / int(self.d_count_dic[i]))  # 分钟平均探测WARN数
            total_y.append(int(self.d_info_dic[i]) + int(self.d_warn_dic[i]))  # 分钟评卷探测 总数
        for i in x:
            # 探测正确率
            right_persent_y.append(float(self.d_info_dic[i]) / float(int(self.d_info_dic[i]) + int(self.d_warn_dic[i])))
        x = self.d_x
        # plt.plot(x, info_y)
        # plt.title("success / day (average)")
        # plt.savefig("charts/image/d_info.jpg")
        # plt.close()
        plt.bar(x, info_y, width=0.5, color='b', lw=100)
        plt.grid(True)
        plt.legend(loc=0)
        plt.axis('tight')
        plt.xlabel('day')
        plt.title("success / day (average)")
        plt.savefig("charts/image/d_info.jpg")
        plt.close()


        # plt.plot(x, warn_y)
        # plt.title("warn / day")
        # plt.savefig("charts/image/d_warn.jpg")
        # plt.close()
        plt.bar(x, warn_y, width=0.5, color='b', lw=100)
        plt.grid(True)
        plt.legend(loc=0)
        plt.axis('tight')
        plt.xlabel('day')
        plt.title("warn / day (average)")
        plt.savefig("charts/image/d_warn.jpg")
        plt.close()

        # plt.plot(x, total_y)
        # plt.title("total / day")
        # plt.savefig("charts/image/d_total.jpg")
        # plt.close()
        plt.bar(x, total_y, width=0.5, color='b', lw=100)
        plt.grid(True)
        plt.legend(loc=0)
        plt.axis('tight')
        plt.xlabel('day')
        plt.title("total / day")
        plt.savefig("charts/image/d_total.jpg")
        plt.close()

        plt.plot(x, right_persent_y)
        plt.title("accuracy rate")
        plt.savefig("charts/image/d_rightP.jpg")
        plt.close()
        pass

    def fuck_end(self):
        """ fuck end !"""
        print "统计文档生成..."
        self.doc = Document()  # 创建 doc 对象
        self.doc.add_heading("Log Analyze", 0)
        self.doc.add_picture("charts/image/m_info.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/m_warn.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/m_total.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/m_rightP.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/h_info.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/h_warn.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/h_total.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/h_rightP.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/d_info.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/d_warn.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/d_total.jpg", width=Inches(5.4))
        self.doc.add_picture("charts/image/d_rightP.jpg", width=Inches(5.4))
        self.doc.save("Analyze_Rst.docx")
        print "完成，结果见 Analyze_Rst.docx"
        pass


if __name__ == '__main__':
    generate_charts().generate_m_charts()
